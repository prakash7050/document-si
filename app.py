import os
import io
import json
import re
from typing import Any

from flask import Flask, request, send_file, jsonify
from docx import Document
from google import genai
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

MAX_FILE_MB = 15
ALLOWED_EXTENSIONS = {".docx"}

# Gemini is only used to understand the template and map JSON keys to
# locations in the document. DOCX editing is done locally with python-docx.
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")


def get_gemini_client() -> genai.Client:
    key = os.getenv("GEMINI_API_KEY")
    if not key:
        raise RuntimeError(
            "GEMINI_API_KEY is not set. Put it in .env or export it in your shell."
        )
    return genai.Client(api_key=key)


def json_path_get(data: Any, path: str) -> Any:
    """
    Supports:
      name
      customer.name
      experiences[0].position
      experiences.0.position
    """
    path = path.strip()
    if not path:
        return None

    # Convert [0] -> .0
    path = re.sub(r"\[(\d+)\]", r".\1", path)

    current = data
    for part in path.split("."):
        if part == "":
            continue
        if isinstance(current, dict):
            if part not in current:
                return None
            current = current[part]
        elif isinstance(current, list) and part.isdigit():
            idx = int(part)
            if idx >= len(current):
                return None
            current = current[idx]
        else:
            return None
    return current


def value_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def replace_paragraph_text(paragraph, value: Any) -> None:
    text = value_to_text(value)

    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell_text(cell, value: Any) -> None:
    """
    Replace text while keeping the table/cell formatting.
    We write into the first paragraph and clear the other runs/paragraphs.
    """
    text = value_to_text(value)

    if not cell.paragraphs:
        cell.add_paragraph()

    first = cell.paragraphs[0]

    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)

    # Clear any additional paragraphs so old placeholder text does not remain.
    for p in cell.paragraphs[1:]:
        for run in p.runs:
            run.text = ""


def iter_all_paragraphs(doc):
    # Body paragraphs
    for i, p in enumerate(doc.paragraphs):
        yield {"type": "paragraph", "index": i, "text": p.text}

    # Paragraphs inside tables are represented through their cells, so the
    # table-cell targets below are enough for this service.


def extract_document_structure(doc: Document) -> dict:
    paragraphs = [
        {"index": i, "text": p.text}
        for i, p in enumerate(doc.paragraphs)
    ]

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                cells.append({
                    "row": ri,
                    "col": ci,
                    "text": cell.text,
                })
            rows.append(cells)
        tables.append({
            "table": ti,
            "rows": rows,
        })

    return {
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_placeholder_keys(doc: Document) -> list[str]:
    """
    If a template contains {{name}}, {{loan_id}}, etc., these are used directly.
    This path needs no AI and is the most deterministic option.
    """
    found = []

    def collect(text: str):
        for key in re.findall(r"\{\{\s*([A-Za-z_][A-Za-z0-9_.\[\]]*)\s*\}\}", text):
            if key not in found:
                found.append(key)

    for p in doc.paragraphs:
        collect(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                collect(cell.text)

    return found


def replace_placeholders(doc: Document, data: dict) -> tuple[list[str], list[str]]:
    filled = []
    missing = []

    def replace_in_paragraph(p):
        original = p.text
        matches = re.findall(
            r"\{\{\s*([A-Za-z_][A-Za-z0-9_.\[\]]*)\s*\}\}", original
        )
        if not matches:
            return

        new_text = original
        for key in matches:
            value = json_path_get(data, key)
            if value is None:
                missing.append(key)
                continue
            new_text = re.sub(
                r"\{\{\s*" + re.escape(key) + r"\s*\}\}",
                value_to_text(value),
                new_text,
            )
            if key not in filled:
                filled.append(key)

        if new_text != original:
            replace_paragraph_text(p, new_text)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text
                matches = re.findall(
                    r"\{\{\s*([A-Za-z_][A-Za-z0-9_.\[\]]*)\s*\}\}", original
                )
                if not matches:
                    continue

                new_text = original
                for key in matches:
                    value = json_path_get(data, key)
                    if value is None:
                        missing.append(key)
                        continue
                    new_text = re.sub(
                        r"\{\{\s*" + re.escape(key) + r"\s*\}\}",
                        value_to_text(value),
                        new_text,
                    )
                    if key not in filled:
                        filled.append(key)

                if new_text != original:
                    replace_cell_text(cell, new_text)

    return sorted(set(filled)), sorted(set(missing))


def normalize_label(value: str) -> str:
    """Normalize a DOCX table label for reliable key matching."""
    value = (value or "").strip().lower()
    value = re.sub(r"[^a-z0-9]+", "_", value).strip("_")
    aliases = {
        "date": "date",
        "dates": "date",
        "employer": "employer",
        "company": "employer",
        "organization": "employer",
        "organisation": "employer",
        "project": "project",
        "role": "role",
        "position": "role",
        "job_title": "role",
        "description": "description",
        "details": "description",
    }
    return aliases.get(value, value)


def fill_repeating_object_tables(
    doc: Document, data: dict
) -> tuple[list[str], set[str]]:
    """
    Deterministically fill repeated DOCX tables from list-of-object JSON fields.

    Example:
      professional_experience = [
        {"date": "...", "employer": "...", "project": "...", ...},
        ...
      ]

    If the DOCX contains repeated tables whose first column has labels
    Date/Employer/Project/Role/Description, each array item is written to
    the next matching table. This avoids relying on an LLM to count repeated
    blocks and makes local and Render behaviour deterministic.
    """
    filled: list[str] = []
    consumed_top_keys: set[str] = set()

    list_fields = [
        (key, value)
        for key, value in data.items()
        if isinstance(value, list)
        and value
        and all(isinstance(item, dict) for item in value)
    ]

    for root_key, records in list_fields:
        candidate_tables = []

        for ti, table in enumerate(doc.tables):
            label_map: dict[str, int] = {}

            for ri, row in enumerate(table.rows):
                if not row.cells:
                    continue
                label = normalize_label(row.cells[0].text)
                if label:
                    label_map[label] = ri

            # Require at least two recognizable field labels before treating
            # a table as a repeated data table. This prevents Revision History
            # and unrelated tables from being selected accidentally.
            record_keys = {normalize_label(k) for k in records[0].keys()}
            overlap = record_keys.intersection(label_map.keys())

            if len(overlap) >= 2:
                candidate_tables.append((ti, label_map, len(overlap)))

        # Use document order so array[0] -> first matching table, etc.
        candidate_tables.sort(key=lambda item: item[0])

        if not candidate_tables:
            continue

        used_count = min(len(records), len(candidate_tables))

        for index in range(used_count):
            ti, label_map, _ = candidate_tables[index]
            table = doc.tables[ti]
            record = records[index]

            for key, value in record.items():
                normalized_key = normalize_label(key)
                row_index = label_map.get(normalized_key)

                if row_index is None:
                    continue

                row = table.rows[row_index]
                if len(row.cells) < 2:
                    continue

                replace_cell_text(row.cells[1], value)

                source = f"{root_key}[{index}].{key}"
                filled.append(source)

        # Only remove this field from the AI workload when every array item
        # received a corresponding table. If the template has fewer blocks,
        # Gemini can still attempt to map the remaining data.
        if used_count == len(records):
            consumed_top_keys.add(root_key)

    return sorted(set(filled)), consumed_top_keys



def fill_cv_identity_paragraphs(
    doc: Document, data: dict
) -> tuple[list[str], set[str]]:
    """
    Fill the common CV title/name/position sequence without AI when present.
    This is a safe deterministic optimization; arbitrary templates still use
    Gemini as the fallback.
    """
    filled: list[str] = []
    consumed: set[str] = set()

    if "name" not in data and "position" not in data:
        return filled, consumed

    # Locate the "Curriculum Vitae" heading and then the next two non-empty
    # body paragraphs. In the supplied DNV template these are Name and Position.
    title_index = None
    for i, p in enumerate(doc.paragraphs):
        if normalize_label(p.text) == "curriculum_vitae":
            title_index = i
            break

    if title_index is None:
        return filled, consumed

    non_empty_after = []
    for i in range(title_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            non_empty_after.append(i)
        if len(non_empty_after) >= 2:
            break

    if "name" in data and len(non_empty_after) >= 1:
        replace_paragraph_text(doc.paragraphs[non_empty_after[0]], data["name"])
        filled.append("name")
        consumed.add("name")

    if "position" in data and len(non_empty_after) >= 2:
        replace_paragraph_text(
            doc.paragraphs[non_empty_after[1]], data["position"]
        )
        filled.append("position")
        consumed.add("position")

    return sorted(set(filled)), consumed


def fill_common_labelled_paragraphs(
    doc: Document, data: dict
) -> tuple[list[str], set[str]]:
    """
    Fill common heading/value paragraph pairs deterministically.
    Gemini remains the fallback for arbitrary/dynamic fields.
    """
    filled: list[str] = []
    consumed: set[str] = set()

    aliases = {
        "summary": {"summary", "profile", "professional_summary"},
        "education": {
            "education",
            "qualifications",
            "education_qualifications",
            "education_qualifications_affiliations",
        },
    }

    for i, paragraph in enumerate(doc.paragraphs):
        label = normalize_label(paragraph.text)

        matched_key = None
        for data_key, possible_labels in aliases.items():
            if label in possible_labels:
                for candidate in possible_labels:
                    if candidate in data:
                        matched_key = candidate
                        break
                if matched_key:
                    break

        if not matched_key:
            continue

        if i + 1 >= len(doc.paragraphs):
            continue

        # Only fill the following paragraph when it is not another heading.
        next_paragraph = doc.paragraphs[i + 1]
        if normalize_label(next_paragraph.text) in {
            "professional_experience",
            "revision_history",
        }:
            continue

        replace_paragraph_text(next_paragraph, data[matched_key])
        filled.append(matched_key)
        consumed.add(matched_key)

    return sorted(set(filled)), consumed


def build_mapping_with_gemini(doc_structure: dict, data: dict) -> dict:
    """
    Ask Gemini to map JSON source paths to actual DOCX locations.
    Returns only executable assignments; document editing stays local.
    """
    client = get_gemini_client()

    # google-genai response_schema does not accept JSON Schema unions like
    # ["integer", "null"]. We use -1 as "not applicable".
    schema = {
        "type": "object",
        "properties": {
            "assignments": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "source": {"type": "string"},
                        "target_type": {
                            "type": "string",
                            "enum": ["paragraph", "table_cell"],
                        },
                        "paragraph_index": {
                            "type": "integer",
                            "description": "Paragraph index; use -1 when not applicable.",
                        },
                        "table_index": {
                            "type": "integer",
                            "description": "Table index; use -1 when not applicable.",
                        },
                        "row": {
                            "type": "integer",
                            "description": "Table row index; use -1 when not applicable.",
                        },
                        "col": {
                            "type": "integer",
                            "description": "Table column index; use -1 when not applicable.",
                        },
                    },
                    "required": [
                        "source",
                        "target_type",
                        "paragraph_index",
                        "table_index",
                        "row",
                        "col",
                    ],
                },
            }
        },
        "required": ["assignments"],
    }

    prompt = f"""
You are a DOCX form-field mapping engine.

Map JSON DATA KEYS to locations in the supplied DOCX template.

RULES:
1. Use ONLY locations that actually exist in DOCUMENT STRUCTURE.
2. Source is a JSON path, e.g. name, customer.name,
   professional_qualifications[0].qualification,
   professional_experience[0].position.
3. For arrays/repeated records, map items in order to corresponding repeated
   blocks in the document.
4. Never invent paragraph/table/row/column indexes.
5. Prefer the blank value cell next to a visible label.
6. A row containing "Qualification:" should map qualification data to its
   associated blank value area.
7. A repeated "Position:" block should receive the corresponding position.
8. A row containing "to" can contain from/start and to/end date values.
9. Do not map unrelated JSON keys to arbitrary locations.
10. paragraph_index, table_index, row and col MUST always be integers.
    Use -1 when a field does not apply.
11. Return ONLY executable assignments. No explanation.

DOCUMENT STRUCTURE:
{json.dumps(doc_structure, ensure_ascii=False, indent=2)}

DATA:
{json.dumps(data, ensure_ascii=False, indent=2)}
"""

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config={
            "response_mime_type": "application/json",
            "response_schema": schema,
        },
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    return json.loads(response.text)


def apply_assignments(doc: Document, assignments: list[dict], data: dict):
    filled_sources = []
    unmatched_sources = []

    # De-duplicate because merged Word cells can appear multiple times in
    # python-docx's row.cells representation.
    seen_targets = set()

    for item in assignments:
        source = item["source"]
        value = json_path_get(data, source)

        if value is None:
            if source not in unmatched_sources:
                unmatched_sources.append(source)
            continue

        target_type = item["target_type"]

        if target_type == "paragraph":
            idx = item["paragraph_index"]
            if idx is None or idx < 0 or idx >= len(doc.paragraphs):
                continue

            target_id = ("p", idx)
            if target_id in seen_targets:
                continue

            replace_paragraph_text(doc.paragraphs[idx], value)
            seen_targets.add(target_id)
            filled_sources.append(source)

        elif target_type == "table_cell":
            ti = item["table_index"]
            ri = item["row"]
            ci = item["col"]

            if ti is None or ri is None or ci is None:
                continue
            if ti < 0 or ti >= len(doc.tables):
                continue

            table = doc.tables[ti]
            if ri < 0 or ri >= len(table.rows):
                continue
            if ci < 0 or ci >= len(table.rows[ri].cells):
                continue

            cell = table.rows[ri].cells[ci]

            # Merged cells can repeat the same underlying XML cell.
            target_id = ("c", id(cell._tc))
            if target_id in seen_targets:
                continue

            replace_cell_text(cell, value)
            seen_targets.add(target_id)
            filled_sources.append(source)

    return sorted(set(filled_sources)), sorted(set(unmatched_sources))


def validate_upload(file_storage):
    if not file_storage or not file_storage.filename:
        raise ValueError("file is required")

    filename = file_storage.filename
    if not filename.lower().endswith(".docx"):
        raise ValueError("Only .docx files are supported")

    # Check size without trusting the extension.
    stream = file_storage.stream
    current = stream.tell()
    stream.seek(0, io.SEEK_END)
    size = stream.tell()
    stream.seek(current)

    if size > MAX_FILE_MB * 1024 * 1024:
        raise ValueError(f"File too large. Maximum is {MAX_FILE_MB} MB.")


@app.get("/health")
def health():
    return jsonify({
        "success": True,
        "service": "DOCX Auto Fill API",
        "ai_provider": "gemini",
        "model": GEMINI_MODEL,
        "gemini_key_configured": bool(os.getenv("GEMINI_API_KEY")),
        "version": "2.0-deterministic-repeat-tables",
    })


@app.get("/")
def home():
    return jsonify({
        "message": "DOCX Auto Fill API is running",
        "endpoints": {
            "POST /analyze": "multipart/form-data: file=<template.docx>",
            "POST /fill-docx": "multipart/form-data: file=<template.docx>, data=<JSON string>",
        },
    })


@app.post("/analyze")
def analyze():
    try:
        file = request.files.get("file")
        validate_upload(file)

        file.stream.seek(0)
        doc = Document(file.stream)
        structure = extract_document_structure(doc)
        placeholders = extract_placeholder_keys(doc)

        return jsonify({
            "success": True,
            "placeholders": placeholders,
            "document": structure,
            "note": (
                "If placeholders are present, use {{field_name}} in the DOCX "
                "for deterministic filling. Otherwise /fill-docx uses Gemini "
                "to map JSON keys to labels/table locations."
            ),
        })

    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)}), 400


@app.post("/fill-docx")
def fill_docx():
    try:
        file = request.files.get("file")
        validate_upload(file)

        raw_data = request.form.get("data")
        if not raw_data:
            return jsonify({
                "success": False,
                "error": "data field is required and must contain JSON"
            }), 400

        try:
            data = json.loads(raw_data)
        except json.JSONDecodeError as exc:
            return jsonify({
                "success": False,
                "error": f"Invalid JSON in data field: {exc}"
            }), 400

        if not isinstance(data, dict):
            return jsonify({
                "success": False,
                "error": "data must be a JSON object"
            }), 400

        file.stream.seek(0)
        doc = Document(file.stream)

        # 1) Deterministic placeholder mode.
        placeholders = extract_placeholder_keys(doc)
        filled, missing = replace_placeholders(doc, data)

        # 2) Deterministic label/table mode.
        #
        # Repeated list-of-object sections (for example 5 professional
        # experience records) are filled directly into matching repeated
        # tables. This is intentionally done BEFORE Gemini so the result is
        # identical on localhost and Render and does not depend on model
        # output/counting.
        deterministic_filled, consumed_keys = fill_repeating_object_tables(
            doc, data
        )
        identity_filled, identity_consumed = fill_cv_identity_paragraphs(
            doc, data
        )
        paragraph_filled, paragraph_consumed = fill_common_labelled_paragraphs(
            doc, data
        )

        filled = sorted(
            set(
                filled
                + deterministic_filled
                + identity_filled
                + paragraph_filled
            )
        )
        consumed_keys.update(identity_consumed)
        consumed_keys.update(paragraph_consumed)

        # 3) Gemini fallback for anything that was not deterministically filled.
        mapping_used = False
        if not placeholders:
            remaining_data = {
                key: value
                for key, value in data.items()
                if key not in consumed_keys
            }

            if remaining_data:
                structure = extract_document_structure(doc)
                mapping = build_mapping_with_gemini(structure, remaining_data)

                ai_filled, ai_unmatched = apply_assignments(
                    doc, mapping.get("assignments", []), remaining_data
                )
                filled = sorted(set(filled + ai_filled))

                mapped_top_keys = {
                    s.split(".")[0].split("[")[0] for s in ai_filled
                }

                remaining_top_keys = list(remaining_data.keys())
                unmatched = [
                    k for k in remaining_top_keys
                    if k not in mapped_top_keys
                ]

                missing = sorted(
                    set(missing + ai_unmatched + unmatched)
                )
                mapping_used = True

        # Any list field that was fully filled deterministically is no longer
        # considered missing.
        missing = sorted(
            set(missing)
            - {
                key
                for key in consumed_keys
            }
        )

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        response = send_file(
            output,
            as_attachment=True,
            download_name="filled_document.docx",
            mimetype=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        response.headers["X-Filled-Fields"] = json.dumps(
            sorted(set(filled)), ensure_ascii=True
        )
        response.headers["X-Unmatched-Fields"] = json.dumps(
            sorted(set(missing)), ensure_ascii=True
        )
        if placeholders:
            mapping_mode = "placeholder"
        elif mapping_used and deterministic_filled:
            mapping_mode = "deterministic+gemini"
        elif deterministic_filled:
            mapping_mode = "deterministic"
        else:
            mapping_mode = "gemini"

        response.headers["X-Mapping-Mode"] = mapping_mode
        return response

    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)}), 500


if __name__ == "__main__":
    # Local + Render fallback. Render provides PORT automatically.
    port = int(os.getenv("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=False)